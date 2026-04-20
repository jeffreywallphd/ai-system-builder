from __future__ import annotations

import html
from dataclasses import dataclass
from pathlib import Path

from ..models import (
    DatasetPreparationSourceInput,
    DatasetPreparationWarning,
    DocumentNormalizationConfig,
)


@dataclass
class NormalizedDocument:
    artifact_id: str
    markdown: str
    media_type: str | None
    source_path: str


@dataclass
class DocumentNormalizationResult:
    documents: list[NormalizedDocument]
    skipped_document_count: int
    warnings: list[DatasetPreparationWarning]


_SUPPORTED_SUFFIXES = {".txt", ".md", ".html", ".htm", ".pdf", ".docx"}

_UNSUPPORTED_BUT_COMMON_SUFFIXES = {".doc"}


def _extension_for_source(source: DatasetPreparationSourceInput) -> str:
    if source.originalName:
        original_extension = Path(source.originalName).suffix.lower()
        if original_extension:
            return original_extension

    return Path(source.localPath).suffix.lower()


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _normalize_html(path: Path) -> str:
    try:
        from markdownify import markdownify as html_to_markdown
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("markdownify is required for HTML normalization") from error

    return html_to_markdown(_read_text(path), heading_style="ATX")


def _normalize_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("pypdf is required for PDF normalization") from error

    reader = PdfReader(str(path))
    page_text: list[str] = []
    for page in reader.pages:
        extracted = page.extract_text() or ""
        if extracted.strip():
            page_text.append(extracted.strip())

    return "\n\n".join(page_text)


def _normalize_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as error:  # pragma: no cover - covered through policy behavior
        raise RuntimeError("python-docx is required for DOCX normalization") from error

    document = Document(str(path))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(lines)


def _normalize_source_to_markdown(source: DatasetPreparationSourceInput) -> str:
    path = Path(source.localPath)
    if not path.exists():
        raise ValueError(f"Input path does not exist: {source.localPath}")

    extension = _extension_for_source(source)
    if extension in _UNSUPPORTED_BUT_COMMON_SUFFIXES:
        raise ValueError(
            "Unsupported document type: .doc (legacy Microsoft Word). Convert to .docx before dataset preparation."
        )

    if extension not in _SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported document type: {extension or 'unknown'}")

    if extension == ".md":
        return _read_text(path)

    if extension == ".txt":
        return _read_text(path)

    if extension in {".html", ".htm"}:
        return _normalize_html(path)

    if extension == ".pdf":
        return _normalize_pdf(path)

    if extension == ".docx":
        return _normalize_docx(path)

    raise ValueError(f"Unsupported document type: {html.escape(extension)}")


def normalize_sources_to_markdown(
    source_inputs: list[DatasetPreparationSourceInput],
    config: DocumentNormalizationConfig,
) -> DocumentNormalizationResult:
    if config.targetFormat != "markdown":
        raise ValueError(f"Unsupported normalization target format: {config.targetFormat}")

    policy = config.unsupportedDocumentPolicy or "fail"
    warnings: list[DatasetPreparationWarning] = []
    normalized: list[NormalizedDocument] = []
    skipped = 0

    for source in source_inputs:
        try:
            markdown = _normalize_source_to_markdown(source)
            normalized.append(
                NormalizedDocument(
                    artifact_id=source.artifactId,
                    markdown=markdown,
                    media_type=source.mediaType,
                    source_path=source.localPath,
                )
            )
        except Exception as error:
            if policy == "skip":
                skipped += 1
                warnings.append(
                    DatasetPreparationWarning(
                        code="document_normalization_skipped",
                        message=(
                            f"Skipped source '{source.artifactId}' during normalization: {error}"
                        ),
                        sourceArtifactId=source.artifactId,
                    )
                )
                continue
            raise

    return DocumentNormalizationResult(
        documents=normalized,
        skipped_document_count=skipped,
        warnings=warnings,
    )
